import logging
from pathlib import Path

import docx
import fitz  # PyMuPDF
import pandas as pd

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class Filehandle:
    """A class to handle reading text content from various file types."""

    def _read_pdf(self, path: Path) -> str:
        """Extracts all text from a PDF file."""
        try:
            text = ""
            with fitz.open(path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Failed to read PDF file {path}: {e}")
            return ""

    def _read_word(self, path: Path) -> str:
        """Extracts text from a DOCX file, including paragraphs and tables."""
        try:
            doc = docx.Document(path)
            full_text = []
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to read DOCX file {path}: {e}")
            return ""

    def _read_excel(self, path: Path) -> str:
        """Extracts text from an Excel file (XLSX/XLS) sheet by sheet."""
        try:
            excel_file = pd.ExcelFile(path)
            text_parts = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                # Convert entire sheet to a string representation
                sheet_text = df.to_string(header=False, index=False, na_rep="")
                text_parts.append(f"--- Sheet: {sheet_name} ---\n{sheet_text}")
            # Join sheets with a clear separator
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to read Excel file {path}: {e}")
            return ""

    def _read_csv(self, path: Path) -> str:
        """Reads a CSV file and returns its content as a formatted string."""
        try:
            df = pd.read_csv(path)
            # to_string provides a nice text-based table representation
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Failed to read CSV file {path}: {e}")
            return ""

    def file_path(self, path_str: str) -> str:
        """
        Reads a file and returns its text content.

        This method acts as a dispatcher, calling the appropriate
        internal reader based on the file's extension.

        Args:
            path_str: The string path to the file.

        Returns:
            The extracted text from the file, or an empty string on error.

        Raises:
            ValueError: If the file format is not supported.
        """
        path = Path(path_str)
        if not path.is_file():
            logger.error(f"File not found at path: {path_str}")
            return ""

        # Map file extensions to their corresponding reader methods
        extension_map = {
            ".pdf": self._read_pdf,
            ".docx": self._read_word,
            ".xlsx": self._read_excel,
            ".xls": self._read_excel,
            ".csv": self._read_csv,
        }

        file_ext = path.suffix.lower()
        reader_func = extension_map.get(file_ext)

        if reader_func:
            return reader_func(path)

        logger.warning(f"Unsupported file format: {file_ext}")
        raise ValueError(f"Unsupported file format: {file_ext}")